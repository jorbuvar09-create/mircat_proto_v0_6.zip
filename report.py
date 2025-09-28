from docx import Document
from docx.shared import Inches
import pandas as pd
import matplotlib.pyplot as plt
from pathlib import Path
import numpy as np

def save_chart(series, title, xlabel, ylabel, outpath):
    plt.figure()
    plt.plot(series.index if hasattr(series, 'index') else range(len(series)),
             series.values if hasattr(series, 'values') else series)
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.tight_layout()
    plt.savefig(outpath, dpi=300, bbox_inches="tight")
    plt.close()

def export_word(prices_df: pd.DataFrame, feat_df: pd.DataFrame, summary: dict, out_dir: str):
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)

    # Charts
    p1 = out/"Fig1_Precio.png"
    p2 = out/"Fig2_VolLocal.png"
    p3 = out/"Fig3_Burst.png"
    p4 = out/"Fig4_Liquidez.png"
    p5 = out/"Fig5_PropCorr.png"
    p6 = out/"Fig6_Wavelet_E8.png"
    p7 = out/"Fig7_Liq_Roll.png"
    p8 = out/"Fig8_Wavelet_Heatmap.png"
    save_chart(prices_df["price"], "Serie de Precios", "t", "Precio", p1)
    save_chart(feat_df["vol_local"].fillna(0), "Volatilidad Local (rolling)", "t", "Vol anualizada", p2)
    save_chart(feat_df["burst"].fillna(0), "Índice de Burst (rolling)", "t", "|Z|", p3)
    save_chart(feat_df["liq_proxy"].fillna(0), "Proxy de Liquidez (rango/price)", "t", "Relativo", p4)
    save_chart(feat_df["prop_corr"].fillna(0), "Correlación Rolling (ret vs lag1)", "t", "ρ", p5)
    # Wavelet energy (ejemplo escala 8)
    if "E_scale8" in feat_df.columns:
        save_chart(feat_df["E_scale8"].fillna(0), "Energía Wavelet (escala 8)", "t", "E", p6)
    # Liquidez Roll proxy
    if "liq_roll" in feat_df.columns:
        save_chart(feat_df["liq_roll"].fillna(0), "Liquidez (Roll proxy)", "t", "Roll", p7)
    # Heatmap wavelet si hay escalas
    cols = [c for c in feat_df.columns if c.startswith("E_scale")]
    if len(cols) >= 3:
        M = np.vstack([feat_df[c].fillna(0).values for c in cols])
        plt.figure()
        plt.imshow(M, aspect='auto')
        plt.title('Heatmap Wavelet (energía por escala)')
        plt.xlabel('Tiempo')
        plt.ylabel('Escalas (ordenadas)')
        plt.tight_layout()
        plt.savefig(p8, dpi=300, bbox_inches='tight')
        plt.close()


    # Word
    doc = Document()
    doc.add_heading("MIRCAT v0.1 — Informe de Microtramos", 0)
    doc.add_paragraph("Este informe resume métricas básicas por microtramos: volatilidad local, burst, liquidez y propagación.")

    # Tabla resumen
    doc.add_heading("1. Resumen", level=1)
    t = doc.add_table(rows=1, cols=2)
    hdr = t.rows[0].cells
    hdr[0].text = "Métrica"
    hdr[1].text = "Valor"
    for k, v in summary.items():
        row = t.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    # Figuras
    doc.add_heading("2. Figuras", level=1)
    doc.add_paragraph("Figura 1. Serie de precios")
    doc.add_picture(str(p1), width=Inches(5.5))
    doc.add_paragraph("Figura 2. Volatilidad local")
    doc.add_picture(str(p2), width=Inches(5.5))
    doc.add_paragraph("Figura 3. Índice de burst")
    doc.add_picture(str(p3), width=Inches(5.5))
    doc.add_paragraph("Figura 4. Proxy de liquidez")
    doc.add_picture(str(p4), width=Inches(5.5))
    doc.add_paragraph("Figura 5. Correlación rolling ret vs lag1")
    doc.add_picture(str(p5), width=Inches(5.5))
    if (out/"Fig6_Wavelet_E8.png").exists():
        doc.add_paragraph("Figura 6. Energía wavelet (escala 8)")
        doc.add_picture(str(p6), width=Inches(5.5))
    if (out/"Fig7_Liq_Roll.png").exists():
        doc.add_paragraph("Figura 7. Liquidez (proxy de Roll)")
        doc.add_picture(str(p7), width=Inches(5.5))
    if (out/"Fig8_Wavelet_Heatmap.png").exists():
        doc.add_paragraph("Figura 8. Heatmap wavelet (energía por escala)")
        doc.add_picture(str(p8), width=Inches(5.5))

    out_doc = out/"MIRCAT_v0_1_Informe.docx"
    doc.save(out_doc)
    return str(out_doc)
